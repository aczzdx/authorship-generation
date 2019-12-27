import pandas as pd
import numpy as np

df = pd.read_csv("data/authors.csv")
first_name = df['First Name']
middle_name = df['Middle Initial(s)']
last_name = df['Last Name']

for tag in ['First Name', 'Middle Initial(s)', 'Last Name']:
    df[tag] = df[tag].apply(lambda name: name.strip() if not pd.isnull(name) else name)

def get_first_name(first_name):
    if '-' in first_name:
        l1 = first_name.split('-')
        l = ''.join(['-'.join(l[0] for l in l1)])
    elif ' ' in first_name:
        l1 = first_name.split(' ')
        l = ''.join([' '.join(l[0] for l in l1)])
    else:
        l = first_name[0].upper()

    return l

def get_middle(middle_initial):
    if pd.isnull(middle_initial):
        return middle_initial
    else:
        result = ''.join([x for x in middle_initial if x.isalpha()])
        if result.istitle():
            return result[0]
        else:
            return result


def get_last_name(last_name):
    if '-' in last_name:
        l1 = last_name.split('-')
        l = ''.join(['-'.join(l[0] for l in l1)])
    elif ' ' in last_name:
        l1 = last_name.split(' ')
        l = ''.join([' '.join(l[0] for l in l1)])
    else:
        l = last_name[0].upper()

    return l

fn=[];
for i in range(len(df)):
    row = df.iloc[i, :]
    l = get_first_name(row['First Name'])
    fn.append(l)
    i=i+1

m=[];
for i in range(len(df)):
    row = df.iloc[i, :]
    l = get_middle(row['Middle Initial(s)'])
    m.append(l)
    i=i+1

ln = [];
for i in range(len(df)):
    row = df.iloc[i, :]
    l = get_last_name(row['Last Name'])
    ln.append(l)
    i = i + 1

list1 = []
for i in range(len(df)):
    dict1 = {'First name': fn[i], 'Middle Initial': m[i], 'Last Name': ln[i]}
    list1.append(dict1)
    i = i + 1


def get_first_name_initial(first_name, test1, s1):
    fn_i = []
    if not len(first_name) == 1:
        for j in range(len(first_name)):
            if '-' in first_name:

                if j in test1:
                    l = first_name[j] + '.'
                else:
                    l = first_name[j]
            else:
                l = first_name[j]
                if ' ' in l:
                    l = s1 + l
            fn_i.append(l)

        fn_i = [x for x in fn_i if x != ' ']
    else:
        fn_i = first_name

    return fn_i


def get_last_name_initial(last_name, test2, s2):
    ln_i = []
    if not len(last_name) == 1:
        for j in range(len(last_name)):
            if '-' in last_name:
                if j in test2:
                    l = last_name[j] + '.'
                else:
                    l = last_name[j]
            else:
                l = last_name[j]
                if ' ' in l:
                    l = s2 + l
            l = ''.join(l.split())
            ln_i.append(l)

        ln_i = [x for x in ln_i if x != ' ']
    else:
        ln_i = last_name

    return ln_i

#%%
print('Eg: First Name: Xiang-Zhen')
g = input("Enter the initials: ")
f = []
f1 = []
test1 = []
for i in range(len(g)):
    if not g[i] == '.':
        f.append(i)
    else:
        f1.append(i)
for j in range(len(f1)):
    l = f1[j] - j - 1
    test1.append(l)

print('Eg: First Name: Jun Soo')
g1 = input("Enter the initials: ")
for i in range(len(g1)):
    if not g1[i].isalpha():
        s1 = g1[i]
    else:
        s1=''

#%%
fn_ii=[]
for i in range(len(df)):
    first_name=list1[i]['First name']
    l=''.join(get_first_name_initial(first_name,test1,s1))
    l1=''.join(l.split( ))
    #l=get_first_name_initial(test,list1,s)
    fn_ii.append(l1)
fn_ii

#%%
print('Eg: Last Name: Baskin-Sommers')
g2 = input("Enter the initials: ")
q = []
q1 = []
test2 = []
for i in range(len(g2)):
    if not g2[i] == '.':
        q.append(i)
    else:
        q1.append(i)
for j in range(len(q1)):
    l = q1[j] - j - 1
    test2.append(l)

print('Eg: Last Name: van Rooij')
g3 = input("Enter the initials: ")
for i in range(len(g3)):
    if not g3[i].isalpha():
        s2 = g3[i]
    else:
        s2=''



#%%
ln_ii = []
for i in range(len(df)):
    last_name = list1[i]['Last Name']
    l = ''.join(get_last_name_initial(last_name, test2, s2))
    # l1=''.join(l.split( ))
    # l=get_last_name_initial(test2,list1,s2)
    ln_ii.append(l)

ln_ii


#def get_initial(fn_ii,m,ln_ii):
#    initial=[]
#    for i in range(len(df)):
#        ret = fn_ii[i]+'.'
#        if not pd.isnull(m[i]):
#            for t in range(len({'Middle Initial':m[i]})):
#                ret += m[i][t]+'.'
#        ret += ln_ii[i]+'.'
#        initial.append(ret)
#    return initial

#l=get_initial(fn_ii,m,ln_ii)

def get_initial(fn,m,ln):
    ret = fn+'.'
    if not pd.isnull(m):
        for t in range(len(m)):
            ret += m[t]+'.'
    ret += ln+'.'
    return ret
l=[]
for i in range(len(fn_ii)):
    initial = get_initial(fn_ii[i],m[i],ln[i])
    l.append(initial)

#%%
def normalize_name(name):
    if not ' ' in name:
        ret=name[0].upper()+name[1:].lower()
    else:
        ret=name
    return ret







#%%判断重复的缩写 num是重复的对应大表的编号
first_name = df['First Name']
middle_name = df['Middle Initial(s)']
last_name = df['Last Name']
duplicated=set()
num=[]
d_initial=[]
for i in range(0,len(l)):
    if l[i] in l[i+1:]:
        duplicated.add(l[i])
test=list(duplicated)
for j in range(len(l)):
    for k in range(len(duplicated)):
        if l[j]==test[k]:
            num.append(j)
            d_initial.append(l[j])
print(num)
print(d_initial)

#%%
fn_find=[]
ln_find=[]
fni_find=[]
lni_find=[]
m_find=[]
new_ln_initial=[]
new_initial=[]
for i in range(len(num)):
    fn_find.append(first_name[num[i]])
    ln_find.append(last_name[num[i]])
    fni_find.append(fn_ii[num[i]])
    lni_find.append(ln_ii[num[i]])
    m_find.append(m[num[i]])
#print(fn_find)
#print(ln_find)
#print(fni_find)
#print(m_find)
#print(lni_find)

#%%1111111判断是不是完全重名
name_find=[]
for i in range(len(fni_find)):
    name_find.append(fn_find[i]+' '+ln_find[i])
print(name_find)

duplicated_name=set()
duplicated_name_num=[]
duplicated_name_num2=[]
duplicated_name_num3=[]

for i in range(0,len(name_find)):
    if name_find[i] in name_find[i+1:]:
        duplicated_name.add(name_find[i])
        duplicated_name_num.append(i)
print(duplicated_name)
print(duplicated_name_num)
for i in range(len(name_find)):
    for j in range(len(duplicated_name)):
        if name_find[i]==list(duplicated_name)[j]:
            duplicated_name_num3.append(i)
duplicated_name_num2 = list(set(duplicated_name_num3).difference(set(duplicated_name_num)))
print(duplicated_name_num2)
print(duplicated_name_num3)
#%%生成重名的人的简称
new_initial_all_part1=[]
new_initial_all_part2=[]
for i in range(len(fni_find)):
    for j in range(len(duplicated_name_num)):
        if i==duplicated_name_num[j]:
            a=fni_find[i]
            b=m_find[i]
            c=normalize_name(lni_find[i])
            test=get_initial(a, b, c)+'1'
            new_initial_all_part1.append(test)
print(new_initial_all_part1)
print(duplicated_name_num)
for i in range(len(fni_find)):
    for j in range(len(duplicated_name_num2)):
        if i==duplicated_name_num2[j]:
            a = fni_find[i]
            b = m_find[i]
            c = normalize_name(lni_find[i])
            test = get_initial(a, b, c) + '2'
            new_initial_all_part2.append(test)
print(new_initial_all_part2)
print(duplicated_name_num2)


#%%222222222判断姓是否重复
num_all=list(range(len(d_initial)))
duplicated_name_not_num = list(set(num_all).difference(set(duplicated_name_num3)))
print(duplicated_name_not_num)
duplicated_name_not=[]
duplicated_last_name_all=[]

for i in range(len(name_find)):
    for j in range(len(duplicated_name_not_num)):
        if i==duplicated_name_not_num[j]:
            duplicated_name_not.append(name_find[i])
            duplicated_last_name_all.append(normalize_name(ln_find[i]))
print(duplicated_name_not)
print(duplicated_last_name_all)

duplicated_ln=set()
duplicated_ln_num=[]
duplicated_ln_num_=[]
for i in range(0,len(duplicated_last_name_all)):
    if duplicated_last_name_all[i] in duplicated_last_name_all[i+1:]:
        duplicated_ln.add(duplicated_last_name_all[i])

print(duplicated_ln)

for i in range(len(duplicated_last_name_all)):
    for j in range(len(duplicated_ln)):
        if duplicated_last_name_all[i]==list(duplicated_ln)[j]:
            duplicated_ln_num.append(i)
print(duplicated_ln_num)
for i in range(len(duplicated_name_not_num)):
    for j in range(len(duplicated_ln_num)):
        if duplicated_ln_num[j]==i:
            duplicated_ln_num_.append(duplicated_name_not_num[i])
print(duplicated_ln_num_)

ln_duplicated=[]
fn_duplicated=[]
for i in range(len(ln_find)):
    for j in range(len(duplicated_ln_num_)):
        if duplicated_ln_num_[j]==i:
            ln_duplicated.append(normalize_name(ln_find[duplicated_ln_num_[j]]))
            fn_duplicated.append(normalize_name(fn_find[duplicated_ln_num_[j]]))

print(fn_duplicated)
print(ln_duplicated)

#%%生成多一位的first name
new_fn_initial=[]
new_initial_fn=[]
for i in range(len(fn_find)):
    for j in range(len(duplicated_ln_num_)):
        if i==duplicated_ln_num_[j]:
            a=normalize_name(fn_find[i])
            new_fn_initial=a[0]+a[1]
            b=m_find[i]
            c=normalize_name(lni_find[i])
            test = get_initial(new_fn_initial, b, c)
            new_initial_fn.append(test)
print(new_initial_fn)
print(duplicated_ln_num_)

#%%33333333剩下的

rem = list(set(range(0,len(fn_find))).difference(set(duplicated_name_num3+duplicated_ln_num_)))
print(rem)


#%%生成多一位的last name
new_ln_initial=[]
new_initial_ln=[]
for i in range(len(ln_find)):
    for j in range(len(rem)):
        if i==rem[j]:
            a=fni_find[i]
            b=m_find[i]
            c=normalize_name(ln_find[i])
            test = get_initial(a, b, c[0]+c[1])
            new_initial_ln.append(test)
print(new_initial_ln)
print(rem)



#%%
first_name=df['First Name']
middle_name=df['Middle Initial(s)']
last_name=df['Last Name']
data={'First Name':first_name,'Middle Initials':middle_name,'Last Name':last_name,'first Initial':fn_ii,'last Initial:':ln_ii,'Initial':l}

pd.options.display.max_rows = None
pd.options.display.max_columns = None
print(pd.DataFrame(data))


#%%
import pandas as pd
df = pd.read_csv("reference.csv")
x=df['Compound Name + highest degree'].tolist()
y1=df['Affiliation 1 Department, Institution'].tolist()
y2=df['Affiliation 2 Department, Institution'].tolist()
y3=df['Affiliation 3 Department, Institution'].tolist()
city1=df['City (e.g. Brisbane)'].tolist()
state1=df['State'].tolist()
country1=df['Country'].tolist()
city2=df['City'].tolist()
state2=df['State.1'].tolist()
country2=df['Country.1'].tolist()
city3=df['City '].tolist()
state3=df['State.2'].tolist()
country3=df['Country.2'].tolist()

all1=[]
for i in range(len(df)):
    if not pd.isnull(y1[i]):
        ret=y1[i]+', '
    if not pd.isnull(city1[i]):
        ret+=city1[i]+', '
    if not pd.isnull(state1[i]):
        ret+=state1[i]+', '
    if not pd.isnull(country1[i]):
        ret+=country1[i]
    all1.append(ret)

all2=[]
for i in range(len(df)):
    if pd.isnull(y2[i]):
        ret=''
    else:
        ret=y2[i]+', '
        if not pd.isnull(city2[i]):
            ret+=city2[i]+', '
        if not pd.isnull(state2[i]):
            ret+=state2[i]+', '
        if not pd.isnull(country2[i]):
            ret+=country2[i]
    all2.append(ret)

all3=[]
for i in range(len(df)):
    if pd.isnull(y3[i]):
        ret=''
    else:
        ret = y3[i] + ', '
        if not pd.isnull(city3[i]):
            ret += city3[i] + ', '
        if not pd.isnull(state3[i]):
            ret += state3[i] + ', '
        if not pd.isnull(country3[i]):
            ret += country3[i]
    all3.append(ret)

i_inform=[]
for i in range(len(all1)):
    i_inform.append(all1[i])
    i_inform.append(all2[i])
    i_inform.append(all3[i])
i_inform = [x for x in i_inform if x != '']

institution=[i_inform[i] for i in range(len(i_inform))]
ins = list(set(institution))
ins.sort(key = institution.index) #运用索引

num=[]
for i in range(len(y1)):
    for j in range(len(ins)):
        if all1[i]==ins[j]:
            l=j+1
    num.append(l)

num1=[]
for i1 in range(len(y2)):
    for j1 in range(len(ins)):
        if not pd.isnull(y2[i1]):
            if all2[i1]==ins[j1]:
                l=j1+1
        else:
            l=''
    num1.append(l)

num2=[]
for i2 in range(len(y3)):
    for j2 in range(len(ins)):
        if not pd.isnull(y3[i2]):
            if all3[i2]==ins[j2]:
                l=j2+1
        else:
            l=''
    num2.append(l)

combine=[x[i]+' '+str(num[i])+' '+str(num1[i])+' '+str(num2[i]) for i in range(len(y1))]

#%%
from docx import Document
from docx.shared import Inches


document = Document()
document.add_heading('Authorlist', 0)
author_text=''
p=document.add_paragraph('')
for i in range(len(combine)):
    p.add_run(x[i]+' ')
    super_text=p.add_run(str(num[i])+' '+str(num1[i])+' '+str(num2[i]))
    super_text.font.superscript=True
    if not i==(len(combine)-1):
        p.add_run(',')

document.add_paragraph('\n')
for j in range(len(ins)):
    document.add_paragraph(ins[j],style='List Number')

document.save('demo.docx')