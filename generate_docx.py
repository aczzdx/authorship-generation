import re

import pandas as pd


class InitialsGenerator:
    """Generate initials of author names

    Extract initials from authors' first name, middle name and last name respectively and then mix them up by some characters like '.', '-'. For some special cases, rules of initials generation can be defined by users. Some examples will be shown on the UI system and users can decide how to extract and combine initials and symbols

    """

    def __init__(self):
        self.last_name_tag = 'Last Name'
        self.middle_initial_tag = 'Middle Initial(s)'
        self.first_name_tag = 'First Name'
        self.first_initial_tag = 'First Initial'
        self.last_initial_tag = 'Last Initial'
        self.initial_tag = 'Initial'
        self.initials_examples = {
            "Xiang-Zhen": "X-Z",
            'Jun Soo': "J-S",
            'Baskin-Sommers': "B-S",
            'van Rooij': "vR"
        }

    def transform(self, df):
        """Transform authors names to initials. Given some initial examples, the form of these initials will be defined by users

        :param df: dataframe
        A dataframe with original authors information
        :return data: dataframe
        An updated dataframe with authors information that adds one column of generated initials
        """

        first_name = df[self.first_name_tag]
        middle_name = df[self.middle_initial_tag]
        last_name = df[self.last_name_tag]

        for tag in [self.first_name_tag, self.middle_initial_tag, self.last_name_tag]:
            df[tag] = df[tag].apply(lambda name: name.strip() if not pd.isnull(name) else name)

        def get_first_name(first_name):
            """Extract initials and special symbols from authors first name in the dataframe

            :param first_name: string
            First name of each author in the dataframe

            :return l: string
            Extracted initials and special symbols
            """
            if '-' in first_name:
                l = ''.join(['-'.join(l[0] for l in first_name.split('-'))])
            elif ' ' in first_name:
                l = ''.join([' '.join(l[0] for l in first_name.split(' '))])
            else:
                l = first_name[0].upper()

            return l

        def get_middle(middle_initial):
            """Extract initials and special symbols from authors middle initials in the dataframe

            :param middle_initial: string
            Middle initials of each author in the dataframe

            :return: result: string
            Extracted initials and special symbols
            """
            if pd.isnull(middle_initial):
                return middle_initial
            else:
                result = ''.join([x for x in middle_initial if x.isalpha()])
                if result.istitle():
                    return result[0]
                else:
                    return result

        def get_last_name(last_name):
            """Extract initials and special symbols from authors last name in the dataframe

            :param last_name: string
            Last name of each author in the dataframe

            :return: l: string
             Extracted initials and special symbols
            """
            if '-' in last_name:
                l = ''.join(['-'.join(l[0] for l in last_name.split('-'))])
            elif ' ' in last_name:
                l = ''.join([' '.join(l[0] for l in last_name.split(' '))])
            else:
                l = last_name[0].upper()

            return l

        fn = []
        for i in range(len(df)):
            row = df.iloc[i, :]
            l = get_first_name(row[self.first_name_tag])
            fn.append(l)
            i = i + 1

        m = []
        for i in range(len(df)):
            row = df.iloc[i, :]
            l = get_middle(row[self.middle_initial_tag])
            m.append(l)
            i = i + 1

        ln = []
        for i in range(len(df)):
            row = df.iloc[i, :]
            l = get_last_name(row[self.last_name_tag])
            ln.append(l)
            i = i + 1

        list1 = []
        for i in range(len(df)):
            dict1 = {'First name': fn[i], 'Middle Initial': m[i], self.last_name_tag: ln[i]}
            list1.append(dict1)
            i = i + 1

        def get_first_name_initial(first_name, test1, s1):
            """Combine initials and symbols of authors first name

            :param first_name: string
            First name of each author in the dataframe
            :param test1: list
            A list of integer representing the locations of symbol '.'
            :param s1: string
            A symbol that replaces the space in first name initials
            :return fn_i: string
            First name initials of each author in the dataframe
            """
            fn_i = []
            if not len(first_name) == 1:
                for j in range(len(first_name)):
                    if '-' in first_name:
                        if j in test1:
                            l = first_name[j] + '.'
                        else:
                            l = first_name[j]
                    else:
                        l = first_name[j]
                        if ' ' in l:
                            l = s1 + l
                    fn_i.append(l)

                fn_i = [x for x in fn_i if x != ' ']
            else:
                fn_i = first_name

            return fn_i

        def get_last_name_initial(last_name, test2, s2):
            """Combine initials and symbols of authors last name

            :param last_name: string
            Last name of each author in the dataframe
            :param test2: list
            A list of integer representing the locations of symbol '.'
            :param s2: string
            A symbol that replaces the space in first name initials
            :return ln_i: string
            Last name initials of each author in the dataframe
            """
            ln_i = []
            if not len(last_name) == 1:
                for j in range(len(last_name)):
                    if '-' in last_name:
                        if j in test2:
                            l = last_name[j] + '.'
                        else:
                            l = last_name[j]
                    else:
                        l = last_name[j]
                        if ' ' in l:
                            l = s2 + l
                    l = ''.join(l.split())
                    ln_i.append(l)

                ln_i = [x for x in ln_i if x != ' ']
            else:
                ln_i = last_name

            return ln_i

        # %%
        # in the widget version, we will specify this by the UI.
        # print('Eg: First Name: Xiang-Zhen')
        # g = input("Enter the initials: ")

        g = self.initials_examples['Xiang-Zhen']
        if len(g) == 0:
            print("ERROR: need initial specified for Xiang-Zhen")
            return None

        f = []
        f1 = []
        test1 = []
        for i in range(len(g)):
            if not g[i] == '.':
                f.append(i)
            else:
                f1.append(i)
        for j in range(len(f1)):
            l = f1[j] - j - 1
            test1.append(l)

        # print('Eg: First Name: Jun Soo')
        # g1 = input("Enter the initials: ")
        g1 = self.initials_examples['Jun Soo']
        if len(g1) == 0:
            print("ERROR: need initial specified for Jun Soo")
            return None

        for i in range(len(g1)):
            if not g1[i].isalpha():
                s1 = g1[i]
            else:
                s1 = ''

        # %%
        fn_ii = []
        for i in range(len(df)):
            first_name = list1[i]['First name']
            l = ''.join(get_first_name_initial(first_name, test1, s1))
            l1 = ''.join(l.split())
            fn_ii.append(l1)

        # %%
        # print('Eg: Last Name: Baskin-Sommers')
        # g2 = input("Enter the initials: ")
        g2 = self.initials_examples['Baskin-Sommers']
        if len(g2) == 0:
            print("ERROR: need initial specified for Baskin-Sommers")
            return None

        q = []
        q1 = []
        test2 = []
        for i in range(len(g2)):
            if not g2[i] == '.':
                q.append(i)
            else:
                q1.append(i)
        for j in range(len(q1)):
            l = q1[j] - j - 1
            test2.append(l)

        # print('Eg: Last Name: van Rooij')
        # g3 = input("Enter the initials: ")
        g3 = self.initials_examples['van Rooij']
        if len(g3) == 0:
            print("ERROR: need initial specified for van Rooij")
            return None

        for i in range(len(g3)):
            if not g3[i].isalpha():
                s2 = g3[i]
            else:
                s2 = ''

        # %%
        ln_ii = []
        for i in range(len(df)):
            last_name = list1[i][self.last_name_tag]
            l = ''.join(get_last_name_initial(last_name, test2, s2))
            ln_ii.append(l)

        def get_initial(fn, m, ln):
            """Get initials of each author

            :param fn: string
            First name initial of each author generated before
            :param m: string
            Middle name initial of each author generated before
            :param ln: string
            Last name initial of each author generated before
            :return ret: string
            Initials of each author
            """
            ret = fn + '.'
            if not pd.isnull(m):
                for t in range(len(m)):
                    ret += m[t] + '.'
            ret += ln + '.'
            return ret

        l = []
        for i in range(len(fn_ii)):
            initial = get_initial(fn_ii[i], m[i], ln_ii[i])
            l.append(initial)

        first_name = df[self.first_name_tag]
        middle_name = df[self.middle_initial_tag]
        last_name = df[self.last_name_tag]
        data1 = {
            "First Name": first_name,
            'Middle Initial(s)': m,
            "Last Name": last_name,
            'First Initial': fn_ii,
            'Last Initial': ln_ii,
            'Initial': l
        }

        def normalize_df(df):
            """Normalize forms of information stored in the dataframe
            """
            fn = df[self.first_name_tag].tolist()
            m = df[self.middle_initial_tag].tolist()
            ln = df[self.last_name_tag].tolist()
            for i in range(len(df)):
                df[self.first_name_tag][i] = fn[i][0].upper() + fn[i][1:].lower()
                df[self.last_name_tag][i] = ln[i][0].upper() + ln[i][1:].lower()
                if not pd.isnull(m[i]):
                    df[self.middle_initial_tag][i] = (''.join(list(filter(str.isalpha, m[i])))).upper()
            return df

        data = normalize_df(pd.DataFrame(data1))

        # %%
        def find_duplication(df, tag):
            """Find duplicated rows in the dataframe based on the features stored in tag

            :param df: dataframe
            :param tag: string
            :return c: dataframe
            Extracted duplicated information
            """
            a = df.drop_duplicates(subset=tag, keep='first')
            b = df.drop_duplicates(subset=tag, keep=False)
            c = a.append(b).drop_duplicates(subset=tag, keep=False)
            return c

        data_duplication = data.loc[
            data[self.initial_tag].isin(find_duplication(data, [self.initial_tag])[self.initial_tag])]

        # %%
        def get_all_duplication_initails(df):
            """Figure out the author whose first name and last name are exactly the same. Add '1' after the initials of the first author and add '2' after the initials of the second author

            :param df: dataframe
            Contains the information of authors who have the same initials
            :return all_duplicated: dataframe
            Contains the information of authors who have the same first name and last name
            :return update: dataframe
            Contains the information of authors who have the same first name and last name with updated initials
            """
            tag = [self.first_name_tag, self.last_name_tag]
            all_duplicated = df.loc[df[self.first_name_tag].isin(find_duplication(df, tag)[self.first_name_tag])]
            a = all_duplicated.drop_duplicates(subset=[self.initial_tag], keep='first')
            b = all_duplicated.drop_duplicates(subset=[self.initial_tag], keep='last')
            a[self.initial_tag] = a[self.initial_tag].map(lambda x: x + '1')
            b[self.initial_tag] = b[self.initial_tag].map(lambda x: x + '2')
            update = pd.concat([a, b])
            return all_duplicated, update

        all_duplication, part1 = get_all_duplication_initails(data_duplication)

        # %%
        flag = data_duplication[self.first_name_tag].isin(all_duplication[self.first_name_tag])
        diff_flag = [not f for f in flag]
        res = data_duplication[diff_flag]

        # %%
        def test_ln_dupliction(df, fn, m, ln):
            """Test whether the generated firstname initials are the same. If the result is true, add one more letter to the first initial

            :param df: dataframe
            :param fn: string
            First name of selected author
            :param m: string
            Middle initials of selected author
            :param ln: string
            Last name of selected author
            :return df: dataframe
            """
            test = df.loc[df[self.initial_tag].isin(find_duplication(df, [self.initial_tag])[self.initial_tag])]
            if not len(test) == 0:
                num = len(test[self.first_initial_tag])
                origin = test[self.first_initial_tag].tolist()
                added = [x[num] for x in fn]
                s = [origin[i] + added[i] for i in range(len(origin))]
                l = []
                for i in range(len(test)):
                    initial = get_initial(s[i], m[i], ln[i])
                    l.append(initial)
                test[self.first_initial_tag] = [a for a in s]
                test[self.initial_tag] = [b for b in l]
                df.loc[test.index, self.initial_tag] = test.loc[test.index, self.initial_tag]
                df.loc[test.index, self.first_initial_tag] = test.loc[test.index, self.first_initial_tag]
            return df

            # %%

        def test_fn_dupliction(df, fn, m, ln):
            """Test whether the generated lastname initials are the same. If the result is true, add one more letter to the last initial

            :param df: dataframe
            :param fn: string
            First name of selected author
            :param m: string
            Middle initials of selected author
            :param ln: string
            Last name of selected author
            :return df: dataframe
            """
            test = df.loc[df[self.initial_tag].isin(find_duplication(df, [self.initial_tag])[self.initial_tag])]
            if not len(test) == 0:
                num = len(test[self.last_initial_tag])
                origin = test[self.last_initial_tag].tolist()
                added = [x[num] for x in ln]
                s = [origin[i] + added[i] for i in range(len(origin))]
                l = []
                for i in range(len(test)):
                    initial = get_initial(fn[i], m[i], s[i])
                    l.append(initial)
                test[self.last_initial_tag] = [a for a in s]
                test[self.initial_tag] = [b for b in l]
                df.loc[test.index, self.initial_tag] = test.loc[test.index, self.initial_tag]
                df.loc[test.index, self.last_initial_tag] = test.loc[test.index, self.last_initial_tag]
            return df

        # %%
        def get_ln_duplication_initials(res):
            """
            Figure out the author whose last name are exactly the same. Add one more letter to its first intiails
            """
            tag = [self.last_name_tag]
            lastname_duplicated = res.loc[res[self.last_name_tag].isin(find_duplication(res, tag)['Last Name'])]
            fn = lastname_duplicated[self.first_name_tag].tolist()
            m = lastname_duplicated[self.middle_initial_tag].tolist()
            ln = lastname_duplicated[self.last_initial_tag].tolist()
            s = [x[0] + x[1] for x in fn]
            l = []
            for i in range(len(lastname_duplicated)):
                initial = get_initial(s[i], m[i], ln[i])
                l.append(initial)
            lastname_duplicated[self.first_initial_tag] = [a for a in s]
            lastname_duplicated[self.initial_tag] = [b for b in l]
            result = test_ln_dupliction(lastname_duplicated, fn, m, ln)
            for i in range(5):
                result = test_ln_dupliction(result, fn, m, ln)
            return result

        part2 = get_ln_duplication_initials(res)
        test1 = part2.loc[part2[self.initial_tag].isin(find_duplication(part2, [self.initial_tag])[self.initial_tag])]

        # %%
        def get_fn_duplication_initials(res, part2):
            """
            Figure out the remaining author whose has duplicated initial. Add one more letter to its last intiail
            """
            flag = res[self.last_name_tag].isin(part2[self.last_name_tag])
            diff_flag = [not f for f in flag]
            firstname_duplicated = res[diff_flag]
            fn = firstname_duplicated[self.first_initial_tag].tolist()
            m = firstname_duplicated[self.middle_initial_tag].tolist()
            ln = firstname_duplicated[self.last_name_tag].tolist()
            s = [x[0] + x[1] for x in ln]
            l = []
            for i in range(len(firstname_duplicated)):
                initial = get_initial(fn[i], m[i], s[i])
                l.append(initial)
            firstname_duplicated[self.last_initial_tag] = [a for a in s]
            firstname_duplicated[self.initial_tag] = [b for b in l]
            result = test_fn_dupliction(firstname_duplicated, fn, m, ln)
            for i in range(5):
                result = test_fn_dupliction(result, fn, m, ln)
            return result

        part3 = get_fn_duplication_initials(res, part2)

        # %%
        update_data = pd.concat([part1, part2, part3], join="inner", axis=0)
        data.loc[update_data.index, :] = update_data.loc[update_data.index, :]

        # print(pd.DataFrame(data))

        # generate_docx(l)

        return data


class DocGenerator:
    """Generate a Doc file that contains a list that contains affiliation information of each author and a list of initials combined with the sequence number of institution that represents which institution each author works for
    """

    def __init__(self):
        self.output_doc_filename = "demo2.docx"
        self.whole_name = 'Compound Name + highest degree'
        self.affiliation_tags = [
            'Affiliation 1 Department, Institution',
            'Affiliation 2 Department, Institution',
            'Affiliation 3 Department, Institution'
        ]
        self.role_tag = 'Role(s)'
        self.roles_priority = [
            'Collected the data',
            'Conceived and designed the analysis',
            'Cohort co-investigator',
            'Contributed data or analysis tools',
            'Performed the analysis',
            'Cohort PI',
            'Read, edited and approved the paper',
            'Wrote the paper',
            'Analyzed the data'
        ]

    def get_indices_of_affiliations(self, df, affiliation_tags):
        """Extract indices and information of affilations in dataframe

        :param df: dataframe
        All authors' information
        :param affiliation_tags: list
        Column name of affiliation in dataframe
        :return affiliation_indices: object
        Indices of affiliations
        :return set_of_affiliation: set
        Affiliation of each author without duplication
        """

        df_affiliations = df[affiliation_tags].T
        list_of_affiliations = pd.concat([df_affiliations[name] for name in df_affiliations.columns], axis=0)

        list_of_affiliations.index = [i for i in range(len(list_of_affiliations))]
        list_of_affiliations = list_of_affiliations.to_list()

        set_of_affiliations = list(set(list_of_affiliations))
        set_of_affiliations.sort(key=list_of_affiliations.index)

        name_to_index = {name: i + 1 for (i, name) in enumerate(set_of_affiliations)}

        affiliation_indices = pd.concat([df[tag].apply(lambda x: name_to_index[x]) for tag in affiliation_tags], axis=1)

        return affiliation_indices, set_of_affiliations

    def generate(self, df, initials, old_way=False):
        """Generate a doc file that contains the author list and initials with sequence number of following list

        :param df: dataframe
        All authors' information
        :param initials: string
        Initials of each author
        """
        x = df[self.whole_name]

        i_inform = []

        if old_way:
            y1 = df[self.affiliation_tags[0]].tolist()
            y2 = df[self.affiliation_tags[1]].tolist()
            y3 = df[self.affiliation_tags[2]].tolist()

            for i in range(len(df)):
                i_inform.append(y1[i])
                if not pd.isnull(y2[i]):
                    i_inform.append(y2[i])
                if not pd.isnull(y3[i]):
                    i_inform.append(y3[i])

            institution = [i_inform[i] for i in range(len(i_inform))]
            ins = list(set(institution))
            ins.sort(key=institution.index)
            num = []
            for i in range(len(y1)):
                for j in range(len(ins)):
                    if y1[i] == ins[j]:
                        initials = j + 1
                num.append(initials)
            num1 = []
            for i1 in range(len(y2)):
                for j1 in range(len(ins)):
                    if not pd.isnull(y2[i1]):
                        if y2[i1] == ins[j1]:
                            initials = j1 + 1
                    else:
                        initials = ''
                num1.append(initials)
            num2 = []
            for i2 in range(len(y3)):
                for j2 in range(len(ins)):
                    if not pd.isnull(y3[i2]):
                        if y3[i2] == ins[j2]:
                            initials = j2 + 1
                    else:
                        initials = ''
                num2.append(initials)
            combine = [x[i] + ' ' + str(num[i]) + ' ' + str(num1[i]) + ' ' + str(num2[i]) for i in range(len(y1))]

            # %%
            from docx import Document
            from docx.shared import Inches
            document = Document()
            document.add_heading('Authorlist', 0)
            author_text = ''
            p = document.add_paragraph('')
            for i in range(len(combine)):
                p.add_run(x[i] + ' ')
                super_text = p.add_run(str(num[i]) + ' ' + str(num1[i]) + ' ' + str(num2[i]))
                if not i == (len(combine) - 1):
                    p.add_run(',')
            document.add_paragraph('\n')
            for j in range(len(ins)):
                document.add_paragraph(ins[j], style='List Number')
            document.save('demo.docx')

        else:
            affiliation_indices, set_of_affiliations = self.get_indices_of_affiliations(df, self.affiliation_tags)
            self.generate_doc_content(df, affiliation_indices, set_of_affiliations, initials)

    def generate_doc_content(self, df, affiliation_indices, set_of_affiliations, initials):
        """Generate the content of the doc file

        :param df: dataframe
        All authors' information
        :param affiliation_indices: object
        Indices of affiliations
        :param set_of_affiliations: set
        Affiliation of each author without duplication
        :param initials: list
        List of initials of all authors
        """
        from docx import Document
        document = Document()
        self.generate_authorlist(affiliation_indices, df, document, set_of_affiliations)
        if self.role_tag is not None:
            self.generate_contribution_list(df, document, initials)

        document.save(self.output_doc_filename)

    def generate_authorlist(self, affiliation_indices, df, document, set_of_affiliations):
        """Add an authorlist in the exsiting doc file

        :param affiliation_indices: object
        Indices of affiliations
        :param df: dataframe
        All authors' information
        :param document:
        Existing doc file that generated before
        :param set_of_affiliations: set
        Affiliation of each author without duplication
        """
        document.add_heading('Authorlist', level=2)
        p = document.add_paragraph('')
        x = df[self.whole_name]
        for i in range(len(affiliation_indices)):
            p.add_run(x[i])
            super_text = p.add_run(", ".join([str(affiliation_indices.loc[i, tag]) for tag in self.affiliation_tags]))
            super_text.font.superscript = True
            if not i == (len(affiliation_indices) - 1):
                p.add_run(', ')
        document.add_paragraph('\n')
        for j in range(len(set_of_affiliations)):
            document.add_paragraph(set_of_affiliations[j], style='List Number')

    def generate_contribution_list(self, df, document, initials):
        """Add a author contribution list in the existing file

        :param df: dataframe
        All authors' information
        :param document:
        Existing doc file that generated before
        :param initials: list
        List of initials of all authors
        """
        flatten = lambda l: [item for sublist in l for item in sublist]

        roles = df[self.role_tag].apply(lambda x: re.split(r"\s*[;]\s*", x))
        roles_set = set(flatten(roles.to_list()))

        initials_group_by_roles = {
            tag: [] for tag in roles_set
        }

        for idx in roles.index:
            initial = initials[idx]
            for tag in roles[idx]:
                initials_group_by_roles[tag].append(initial)

        document.add_heading("Author Contributions (in alphabetical order)", level=2)

        role_set_order = self.roles_priority.copy()
        role_set_order.extend(roles_set - set(role_set_order))

        for role_name in roles_set:
            paragraph = document.add_paragraph("")
            title = paragraph.add_run(role_name)
            title.italic = True
            paragraph.add_run("\n")

            author_list = initials_group_by_roles[role_name].copy()
            author_list.sort()

            for i, author in enumerate(author_list):
                if i > 0:
                    paragraph.add_run(", ")

                paragraph.add_run(author)


