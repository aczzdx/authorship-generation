import pandas as pd
import numpy as np

df = pd.read_csv("authors.csv")
first_name = df['First Name']
middle_name = df['Middle Initial(s)']
last_name = df['Last Name']

for tag in ['First Name', 'Middle Initial(s)', 'Last Name']:
    df[tag] = df[tag].apply(lambda name: name.strip() if not pd.isnull(name) else name)

def get_first_name(first_name):
    if '-' in first_name:
        l1 = first_name.split('-')
        l = ''.join(['-'.join(l[0] for l in l1)])
    elif ' ' in first_name:
        l1 = first_name.split(' ')
        l = ''.join([' '.join(l[0] for l in l1)])
    else:
        l = first_name[0].upper()

    return l
/Users/apple/PycharmProjects/authorship-generation/scratch.py

def get_middle(middle_initial):
    if pd.isnull(middle_initial):
        return middle_initial
    else:
        result = ''.join([x for x in middle_initial if x.isalpha()])
        if result.istitle():
            return result[0]
        else:
            return result


def get_last_name(last_name):
    if '-' in last_name:
        l1 = last_name.split('-')
        l = ''.join(['-'.join(l[0] for l in l1)])
    elif ' ' in last_name:
        l1 = last_name.split(' ')
        l = ''.join([' '.join(l[0] for l in l1)])
    else:
        l = last_name[0].upper()

    return l

fn=[];
for i in range(len(df)):
    row = df.iloc[i, :]
    l = get_first_name(row['First Name'])
    fn.append(l)
    i=i+1

m=[];
for i in range(len(df)):
    row = df.iloc[i, :]
    l = get_middle(row['Middle Initial(s)'])
    m.append(l)
    i=i+1

ln = [];
for i in range(len(df)):
    row = df.iloc[i, :]
    l = get_last_name(row['Last Name'])
    ln.append(l)
    i = i + 1

list1 = []
for i in range(len(df)):
    dict1 = {'First name': fn[i], 'Middle Initial': m[i], 'Last Name': ln[i]}
    list1.append(dict1)
    i = i + 1


def get_first_name_initial(first_name, test1, s1):
    fn_i = []
    if not len(first_name) == 1:
        for j in range(len(first_name)):
            if '-' in first_name:

                if j in test1:
                    l = first_name[j] + '.'
                else:
                    l = first_name[j]
            else:
                l = first_name[j]
                if ' ' in l:
                    l = s1 + l
            fn_i.append(l)

        fn_i = [x for x in fn_i if x != ' ']
    else:
        fn_i = first_name

    return fn_i


def get_last_name_initial(last_name, test2, s2):
    ln_i = []
    if not len(last_name) == 1:
        for j in range(len(last_name)):
            if '-' in last_name:
                if j in test2:
                    l = last_name[j] + '.'
                else:
                    l = last_name[j]
            else:
                l = last_name[j]
                if ' ' in l:
                    l = s2 + l
            l = ''.join(l.split())
            ln_i.append(l)

        ln_i = [x for x in ln_i if x != ' ']
    else:
        ln_i = last_name

    return ln_i

#%%
print('Eg: First Name: Xiang-Zhen')
g = input("Enter the initials: ")
f = []
f1 = []
test1 = []
for i in range(len(g)):
    if not g[i] == '.':
        f.append(i)
    else:
        f1.append(i)
for j in range(len(f1)):
    l = f1[j] - j - 1
    test1.append(l)

print('Eg: First Name: Jun Soo')
g1 = input("Enter the initials: ")
for i in range(len(g1)):
    if not g1[i].isalpha():
        s1 = g1[i]
    else:
        s1=''

#%%
fn_ii=[]
for i in range(len(df)):
    first_name=list1[i]['First name']
    l=''.join(get_first_name_initial(first_name,test1,s1))
    l1=''.join(l.split( ))
    #l=get_first_name_initial(test,list1,s)
    fn_ii.append(l1)
fn_ii

#%%
print('Eg: Last Name: Baskin-Sommers')
g2 = input("Enter the initials: ")
q = []
q1 = []
test2 = []
for i in range(len(g2)):
    if not g2[i] == '.':
        q.append(i)
    else:
        q1.append(i)
for j in range(len(q1)):
    l = q1[j] - j - 1
    test2.append(l)

print('Eg: Last Name: van Rooij')
g3 = input("Enter the initials: ")
for i in range(len(g3)):
    if not g3[i].isalpha():
        s2 = g3[i]
    else:
        s2=''



#%%
ln_ii = []
for i in range(len(df)):
    last_name = list1[i]['Last Name']
    l = ''.join(get_last_name_initial(last_name, test2, s2))
    # l1=''.join(l.split( ))
    # l=get_last_name_initial(test2,list1,s2)
    ln_ii.append(l)

ln_ii


def get_initial(fn_ii,m,ln_ii):
    initial=[]
    for i in range(len(df)):
        ret = fn_ii[i]+'.'
        if not pd.isnull(m[i]):
            for t in range(len({'Middle Initial':m[i]})):
                ret += m[i][t]+'.'
        ret += ln_ii[i]+'.'
        initial.append(ret)
    return initial

l=get_initial(fn_ii,m,ln_ii)

first_name=df['First Name']
middle_name=df['Middle Initial(s)']
last_name=df['Last Name']
data={'First Name':first_name,'Middle Initials':middle_name,'Last Name':last_name,'first Initial':fn_ii,'last Initial:':ln_ii,'Initial':l}

pd.options.display.max_rows = None
pd.options.display.max_columns = None
print(pd.DataFrame(data))



#%%
import pandas as pd
df = pd.read_csv("reference.csv")
x=df['Compound Name + highest degree']
y1=df['Affiliation 1 Department, Institution'].tolist()
y2=df['Affiliation 2 Department, Institution'].tolist()
y3=df['Affiliation 3 Department, Institution'].tolist()
i_inform=[]
for i in range(len(df)):
    i_inform.append(y1[i])
    if not pd.isnull(y2[i]):
        i_inform.append(y2[i])
    if not pd.isnull(y3[i]):
        i_inform.append(y3[i])

institution=[i_inform[i] for i in range(len(i_inform))]
ins = list(set(institution))
ins.sort(key = institution.index) #运用索引

num=[]
for i in range(len(y1)):
    for j in range(len(ins)):
        if y1[i]==ins[j]:
            l=j+1
    num.append(l)

num1=[]
for i1 in range(len(y2)):
    for j1 in range(len(ins)):
        if not pd.isnull(y2[i1]):
            if y2[i1]==ins[j1]:
                l=j1+1
        else:
            l=''
    num1.append(l)

num2=[]
for i2 in range(len(y3)):
    for j2 in range(len(ins)):
        if not pd.isnull(y3[i2]):
            if y3[i2]==ins[j2]:
                l=j2+1
        else:
            l=''
    num2.append(l)

combine=[x[i]+' '+str(num[i])+' '+str(num1[i])+' '+str(num2[i]) for i in range(len(y1))]

#%%
from docx import Document
from docx.shared import Inches


document = Document()
document.add_heading('Authorlist', 0)
author_text=''
p=document.add_paragraph('')
for i in range(len(combine)):
    p.add_run(x[i]+' ')
    super_text=p.add_run(str(num[i])+' '+str(num1[i])+' '+str(num2[i]))
    super_text.font.superscript=True
    if not i==(len(combine)-1):
        p.add_run(',')

document.add_paragraph('\n')
for j in range(len(ins)):
    document.add_paragraph(ins[j],style='List Number')

document.save('demo.docx')