import pandas as pd
import re


class InitialsGenerator:

    def __init__(self):
        self.last_name_tag = 'Last Name'
        self.middle_initial_tag = 'Middle Initial(s)'
        self.first_name_tag = 'First Name'
        self.initials_examples = {
            "Xiang-Zhen": "X-Z",
            'Jun Soo': "J-S",
            'Baskin-Sommers': "B-S",
            'van Rooij': "vR"
        }

    def transform(self, df):

        first_name = df[self.first_name_tag]
        middle_name = df[self.middle_initial_tag]
        last_name = df[self.last_name_tag]

        for tag in [self.first_name_tag, self.middle_initial_tag, self.last_name_tag]:
            df[tag] = df[tag].apply(lambda name: name.strip() if not pd.isnull(name) else name)

        def get_first_name(first_name):
            if '-' in first_name:
                l1 = first_name.split('-')
                l = ''.join(['-'.join(l[0] for l in l1)])
            elif ' ' in first_name:
                l1 = first_name.split(' ')
                l = ''.join([' '.join(l[0] for l in l1)])
            else:
                l = first_name[0].upper()

            return l

        def get_middle(middle_initial):
            if pd.isnull(middle_initial):
                return middle_initial
            else:
                result = ''.join([x for x in middle_initial if x.isalpha()])
                if result.istitle():
                    return result[0]
                else:
                    return result


        def get_last_name(last_name):
            if '-' in last_name:
                l1 = last_name.split('-')
                l = ''.join(['-'.join(l[0] for l in l1)])
            elif ' ' in last_name:
                l1 = last_name.split(' ')
                l = ''.join([' '.join(l[0] for l in l1)])
            else:
                l = last_name[0].upper()

            return l

        fn=[]
        for i in range(len(df)):
            row = df.iloc[i, :]
            l = get_first_name(row[self.first_name_tag])
            fn.append(l)
            i=i+1

        m=[]
        for i in range(len(df)):
            row = df.iloc[i, :]
            l = get_middle(row[self.middle_initial_tag])
            m.append(l)
            i=i+1

        ln = []
        for i in range(len(df)):
            row = df.iloc[i, :]
            l = get_last_name(row[self.last_name_tag])
            ln.append(l)
            i = i + 1

        list1 = []
        for i in range(len(df)):
            dict1 = {'First name': fn[i], 'Middle Initial': m[i], self.last_name_tag: ln[i]}
            list1.append(dict1)
            i = i + 1


        def get_first_name_initial(first_name, test1, s1):
            fn_i = []
            if not len(first_name) == 1:
                for j in range(len(first_name)):
                    if '-' in first_name:

                        if j in test1:
                            l = first_name[j] + '.'
                        else:
                            l = first_name[j]
                    else:
                        l = first_name[j]
                        if ' ' in l:
                            l = s1 + l
                    fn_i.append(l)

                fn_i = [x for x in fn_i if x != ' ']
            else:
                fn_i = first_name

            return fn_i


        def get_last_name_initial(last_name, test2, s2):
            ln_i = []
            if not len(last_name) == 1:
                for j in range(len(last_name)):
                    if '-' in last_name:
                        if j in test2:
                            l = last_name[j] + '.'
                        else:
                            l = last_name[j]
                    else:
                        l = last_name[j]
                        if ' ' in l:
                            l = s2 + l
                    l = ''.join(l.split())
                    ln_i.append(l)

                ln_i = [x for x in ln_i if x != ' ']
            else:
                ln_i = last_name

            return ln_i

        #%%
        # in the widget version, we will specify this by the UI.
        # print('Eg: First Name: Xiang-Zhen')
        # g = input("Enter the initials: ")

        g = self.initials_examples['Xiang-Zhen']
        if len(g) == 0:
            print("ERROR: need initial specified for Xiang-Zhen")
            return None

        f = []
        f1 = []
        test1 = []
        for i in range(len(g)):
            if not g[i] == '.':
                f.append(i)
            else:
                f1.append(i)
        for j in range(len(f1)):
            l = f1[j] - j - 1
            test1.append(l)

        # print('Eg: First Name: Jun Soo')
        # g1 = input("Enter the initials: ")
        g1 = self.initials_examples['Jun Soo']
        if len(g1) == 0:
            print("ERROR: need initial specified for Jun Soo")
            return None

        for i in range(len(g1)):
            if not g1[i].isalpha():
                s1 = g1[i]
            else:
                s1=''

        #%%
        fn_ii=[]
        for i in range(len(df)):
            first_name=list1[i]['First name']
            l=''.join(get_first_name_initial(first_name,test1,s1))
            l1=''.join(l.split( ))
            #l=get_first_name_initial(test,list1,s)
            fn_ii.append(l1)
        print(fn_ii)


        #%%
        # print('Eg: Last Name: Baskin-Sommers')
        # g2 = input("Enter the initials: ")
        g2 = self.initials_examples['Baskin-Sommers']
        if len(g2) == 0:
            print("ERROR: need initial specified for Baskin-Sommers")
            return None

        q = []
        q1 = []
        test2 = []
        for i in range(len(g2)):
            if not g2[i] == '.':
                q.append(i)
            else:
                q1.append(i)
        for j in range(len(q1)):
            l = q1[j] - j - 1
            test2.append(l)

        # print('Eg: Last Name: van Rooij')
        # g3 = input("Enter the initials: ")
        g3 = self.initials_examples['van Rooij']
        if len(g3) == 0:
            print("ERROR: need initial specified for van Rooij")
            return None

        for i in range(len(g3)):
            if not g3[i].isalpha():
                s2 = g3[i]
            else:
                s2=''



        #%%
        ln_ii = []
        for i in range(len(df)):
            last_name = list1[i][self.last_name_tag]
            l = ''.join(get_last_name_initial(last_name, test2, s2))
            # l1=''.join(l.split( ))
            # l=get_last_name_initial(test2,list1,s2)
            ln_ii.append(l)

        # ln_ii

        def get_initial(fn, m, ln):
            ret = fn + '.'
            if not pd.isnull(m):
                for t in range(len(m)):
                    ret += m[t] + '.'
            ret += ln + '.'
            return ret

        l = []
        for i in range(len(fn_ii)):
            initial = get_initial(fn_ii[i], m[i], ln_ii[i])
            l.append(initial)

        #l=get_initial(fn_ii,m,ln_ii)

        def normalize_name(name):
            if not ' ' in name:
                ret = name[0].upper() + name[1:].lower()
            else:
                ret = name
            return ret

        first_name = df[self.first_name_tag]
        middle_name = df[self.middle_initial_tag]
        last_name = df[self.last_name_tag]

        def find_duplicate(l):
            duplicated = set()
            num = []
            d_initial = []
            for i in range(0, len(l)):
                if l[i] in l[i + 1:]:
                    duplicated.add(l[i])
            test = list(duplicated)
            for j in range(len(l)):
                for k in range(len(duplicated)):
                    if l[j] == test[k]:
                        num.append(j)
                        d_initial.append(l[j])
            return num, d_initial

        num, d_initial = find_duplicate(l)
        print(num)
        # print(d_initial)

        #%%
        fn_find = []
        ln_find = []
        fni_find = []
        lni_find = []
        m_find = []
        new_ln_initial = []
        new_initial = []
        for i in range(len(num)):
            fn_find.append(first_name[num[i]])
            ln_find.append(last_name[num[i]])
            fni_find.append(fn_ii[num[i]])
            lni_find.append(ln_ii[num[i]])
            m_find.append(m[num[i]])
        #print(fn_find)
        #print(ln_find)
        # print(fni_find)
        # print(m_find)
        # print(lni_find)

        def full_name_find(fn_find, m_find, ln_find):
            name_find = []
            for i in range(len(fn_find)):
                temp_list = [fn_find[i], m_find[i], ln_find[i]]
                temp_list = [x for x in temp_list if type(x) is str]
                name_find.append(" ".join(temp_list))
            return name_find

        name_find = full_name_find(fn_find, m_find, ln_find)
        #print(name_find)

        def all_same_classification(name_find):
            duplicated_name = set()
            duplicated_name_num = []
            duplicated_name_num2 = []
            duplicated_name_num3 = []
            for i in range(0, len(name_find)):
                if name_find[i] in name_find[i + 1:]:
                    duplicated_name.add(name_find[i])
                    duplicated_name_num.append(i)
            for i in range(len(name_find)):
                for j in range(len(duplicated_name)):
                    if name_find[i] == list(duplicated_name)[j]:
                        duplicated_name_num3.append(i)
            duplicated_name_num2 = list(set(duplicated_name_num3).difference(set(duplicated_name_num)))
            return duplicated_name, duplicated_name_num, duplicated_name_num2, duplicated_name_num3

        duplicated_name, duplicated_name_num, duplicated_name_num2, duplicated_name_num3 = all_same_classification(
            name_find)

        #print(duplicated_name)
        #print(duplicated_name_num)
        #print(duplicated_name_num2)
        #print(duplicated_name_num3)

        # %%生成完全重名的人的简称

        def all_same_initials(fni_find, m_find, lni_find):
            new_initial_all_part1 = []
            new_initial_all_part2 = []
            for i in range(len(fni_find)):
                for j in range(len(duplicated_name_num)):
                    if i == duplicated_name_num[j]:
                        a = fni_find[i]
                        b = m_find[i]
                        c = normalize_name(lni_find[i])
                        test = get_initial(a, b, c) + '1'
                        new_initial_all_part1.append(test)

            for i in range(len(fni_find)):
                for j in range(len(duplicated_name_num2)):
                    if i == duplicated_name_num2[j]:
                        a = fni_find[i]
                        b = m_find[i]
                        c = normalize_name(lni_find[i])
                        test = get_initial(a, b, c) + '2'
                        new_initial_all_part2.append(test)
            return new_initial_all_part1, new_initial_all_part2

        new_initial_all_part1, new_initial_all_part2 = all_same_initials(fni_find, m_find, lni_find)
        # print(new_initial_all_part1)
        # print(duplicated_name_num)
        # print(new_initial_all_part2)
        # print(duplicated_name_num2)
        # %%222222222判断姓是否重复
        num_all = list(range(len(d_initial)))
        duplicated_name_not_num = list(set(num_all).difference(set(duplicated_name_num3)))

        def find_duplicated_last_name(duplicated_name_not_num, name_find, ln_find):
            duplicated_name_not = []
            duplicated_last_name_all = []
            for i in range(len(name_find)):
                for j in range(len(duplicated_name_not_num)):
                    if i == duplicated_name_not_num[j]:
                        duplicated_name_not.append(name_find[i])
                        duplicated_last_name_all.append(normalize_name(ln_find[i]))
            duplicated_ln = set()
            duplicated_ln_num = []
            duplicated_ln_num_ = []
            for i in range(0, len(duplicated_last_name_all)):
                if duplicated_last_name_all[i] in duplicated_last_name_all[i + 1:]:
                    duplicated_ln.add(duplicated_last_name_all[i])
            for i in range(len(duplicated_last_name_all)):
                for j in range(len(duplicated_ln)):
                    if duplicated_last_name_all[i] == list(duplicated_ln)[j]:
                        duplicated_ln_num.append(i)
            for i in range(len(duplicated_name_not_num)):
                for j in range(len(duplicated_ln_num)):
                    if duplicated_ln_num[j] == i:
                        duplicated_ln_num_.append(duplicated_name_not_num[i])
            # print(duplicated_ln_num_)
            ln_duplicated = []
            fn_duplicated = []
            for i in range(len(ln_find)):
                for j in range(len(duplicated_ln_num_)):
                    if duplicated_ln_num_[j] == i:
                        ln_duplicated.append(normalize_name(ln_find[duplicated_ln_num_[j]]))
                        fn_duplicated.append(normalize_name(fn_find[duplicated_ln_num_[j]]))
            # print(fn_duplicated)
            # print(ln_duplicated)
            return duplicated_ln_num_, fn_duplicated, ln_duplicated

        duplicated_ln_num_, fn_duplicated, ln_duplicated = find_duplicated_last_name(duplicated_name_not_num, name_find,
                                                                                     ln_find)
        # print(duplicated_ln_num_)
        # print(fn_duplicated)
        # print(ln_duplicated)

        # %%生成多一位的first name
        def duplicated_first_name_initial_generation(fn_find, duplicated_ln_num_, m_find, lni_find):
            new_fn_initial = []
            new_initial_fn = []
            for i in range(len(fn_find)):
                for j in range(len(duplicated_ln_num_)):
                    if i == duplicated_ln_num_[j]:
                        a = normalize_name(fn_find[i])
                        new_fn_initial = a[0] + a[1]
                        b = m_find[i]
                        c = normalize_name(lni_find[i])
                        test = get_initial(new_fn_initial, b, c)
                        new_initial_fn.append(test)
            return new_initial_fn

        new_initial_fn = duplicated_first_name_initial_generation(fn_find, duplicated_ln_num_, m_find, lni_find)
        # print(new_initial_fn)
        # print(duplicated_ln_num_)

        # %%33333333剩下的

        rem = list(set(range(0, len(fn_find))).difference(set(duplicated_name_num3 + duplicated_ln_num_)))

        # %%生成多一位的last name
        def duplicated_last_name_initial_generation(fni_find, m_find, ln_find, rem):
            new_ln_initial = []
            new_initial_ln = []
            for i in range(len(ln_find)):
                for j in range(len(rem)):
                    if i == rem[j]:
                        a = fni_find[i]
                        b = m_find[i]
                        c = normalize_name(ln_find[i])
                        test = get_initial(a, b, c[0] + c[1])
                        new_initial_ln.append(test)
            return new_initial_ln

        new_initial_ln = duplicated_last_name_initial_generation(fni_find, m_find, ln_find, rem)
        # print(new_initial_ln)
        # print(rem)

        # %%
        def num_link(duplicated_name_num, duplicated_name_num2, duplicated_ln_num_, rem, num):
            num1 = []
            num2 = []
            num3 = []
            num4 = []
            for i in range(len(num)):
                for j in range(len(duplicated_name_num)):
                    if i == duplicated_name_num[j]:
                        num1.append(num[i])
            for i in range(len(num)):
                for j in range(len(duplicated_name_num2)):
                    if i == duplicated_name_num2[j]:
                        num2.append(num[i])
            for i in range(len(num)):
                for j in range(len(duplicated_ln_num_)):
                    if i == duplicated_ln_num_[j]:
                        num3.append(num[i])
            for i in range(len(num)):
                for j in range(len(rem)):
                    if i == rem[j]:
                        num4.append(num[i])
            return num1, num2, num3, num4

        num1, num2, num3, num4 = num_link(duplicated_name_num, duplicated_name_num2, duplicated_ln_num_, rem, num)
        # print(num1, num2, num3, num4)
        # print(new_initial_all_part1, new_initial_all_part2, new_initial_fn, new_initial_ln)

        # %%
        def updated_initials(l, num1, num2, num3, num4, new_initial_all_part1, new_initial_all_part2, new_initial_fn,
                             new_initial_ln):
            num = num1 + num2 + num3 + num4
            l_updated = new_initial_all_part1 + new_initial_all_part2 + new_initial_fn + new_initial_ln
            for i in range(len(num)):
                l[num[i]] = l_updated[i]
            return l

        updated_l = updated_initials(l, num1, num2, num3, num4, new_initial_all_part1, new_initial_all_part2,
                                     new_initial_fn, new_initial_ln)
        # print(updated_l)


        first_name=df[self.first_name_tag]
        middle_name=df[self.middle_initial_tag]
        last_name=df[self.last_name_tag]
        data={
            "First Name": first_name,
            'Middle Initials': middle_name,
            "Last Name" :last_name,
            'first Initial':fn_ii,
            'last Initial:':ln_ii,
            'Initial':updated_l
        }

        pd.options.display.max_rows = None
        pd.options.display.max_columns = None


        #print(pd.DataFrame(data))

        # generate_docx(l)

        return pd.DataFrame(data)


class DocGenerator:


    def __init__(self):
        self.output_doc_filename = "demo2.docx"
        self.whole_name = 'Compound Name + highest degree'
        self.affiliation_tags = [
            'Affiliation 1 Department, Institution',
            'Affiliation 2 Department, Institution',
            'Affiliation 3 Department, Institution'
        ]
        self.role_tag = 'Role(s)'
        self.roles_priority = [
            'Collected the data',
            'Conceived and designed the analysis',
            'Cohort co-investigator',
            'Contributed data or analysis tools',
            'Performed the analysis',
            'Cohort PI',
            'Read, edited and approved the paper',
            'Wrote the paper',
            'Analyzed the data'
        ]

    def get_indices_of_affiliations(self, df, affiliation_tags):

        df_affiliations = df[affiliation_tags].T
        list_of_affiliations = pd.concat([df_affiliations[name] for name in df_affiliations.columns], axis=0)

        list_of_affiliations.index = [i for i in range(len(list_of_affiliations))]
        list_of_affiliations = list_of_affiliations.to_list()

        set_of_affiliations = list(set(list_of_affiliations))
        set_of_affiliations.sort(key=list_of_affiliations.index)

        name_to_index = {name: i + 1 for (i, name) in enumerate(set_of_affiliations)}

        affiliation_indices = pd.concat([df[tag].apply(lambda x: name_to_index[x]) for tag in affiliation_tags], axis=1)

        return affiliation_indices, set_of_affiliations

    def generate(self, df, initials, old_way=False):
        # %%
        x = df[self.whole_name]

        i_inform = []

        if old_way:
            y1 = df[self.affiliation_tags[0]].tolist()
            y2 = df[self.affiliation_tags[1]].tolist()
            y3 = df[self.affiliation_tags[2]].tolist()

            for i in range(len(df)):
                i_inform.append(y1[i])
                if not pd.isnull(y2[i]):
                    i_inform.append(y2[i])
                if not pd.isnull(y3[i]):
                    i_inform.append(y3[i])

            institution = [i_inform[i] for i in range(len(i_inform))]
            ins = list(set(institution))
            ins.sort(key=institution.index)  # 运用索引
            num = []
            for i in range(len(y1)):
                for j in range(len(ins)):
                    if y1[i] == ins[j]:
                        initials = j + 1
                num.append(initials)
            num1 = []
            for i1 in range(len(y2)):
                for j1 in range(len(ins)):
                    if not pd.isnull(y2[i1]):
                        if y2[i1] == ins[j1]:
                            initials = j1 + 1
                    else:
                        initials = ''
                num1.append(initials)
            num2 = []
            for i2 in range(len(y3)):
                for j2 in range(len(ins)):
                    if not pd.isnull(y3[i2]):
                        if y3[i2] == ins[j2]:
                            initials = j2 + 1
                    else:
                        initials = ''
                num2.append(initials)
            combine = [x[i] + ' ' + str(num[i]) + ' ' + str(num1[i]) + ' ' + str(num2[i]) for i in range(len(y1))]

            # %%
            from docx import Document
            from docx.shared import Inches
            document = Document()
            document.add_heading('Authorlist', 0)
            author_text = ''
            p = document.add_paragraph('')
            for i in range(len(combine)):
                p.add_run(x[i] + ' ')
                super_text = p.add_run(str(num[i]) + ' ' + str(num1[i]) + ' ' + str(num2[i]))
                if not i == (len(combine) - 1):
                    p.add_run(',')
            document.add_paragraph('\n')
            for j in range(len(ins)):
                document.add_paragraph(ins[j], style='List Number')
            document.save('demo.docx')

        else:
            affiliation_indices, set_of_affiliations = self.get_indices_of_affiliations(df, self.affiliation_tags)
            self.generate_doc_content(df, affiliation_indices, set_of_affiliations, initials)

    def generate_doc_content(self, df, affiliation_indices, set_of_affiliations, initials):
        from docx import Document
        document = Document()
        self.generate_authorlist(affiliation_indices, df, document, set_of_affiliations)
        if self.role_tag is not None:
            self.generate_contribution_list(df, document, initials)

        document.save(self.output_doc_filename)

    def generate_authorlist(self, affiliation_indices, df, document, set_of_affiliations):
        document.add_heading('Authorlist', level=2)
        p = document.add_paragraph('')
        x = df[self.whole_name]
        for i in range(len(affiliation_indices)):
            p.add_run(x[i])
            super_text = p.add_run(", ".join([str(affiliation_indices.loc[i, tag]) for tag in self.affiliation_tags]))
            super_text.font.superscript = True
            if not i == (len(affiliation_indices) - 1):
                p.add_run(', ')
        document.add_paragraph('\n')
        for j in range(len(set_of_affiliations)):
            document.add_paragraph(set_of_affiliations[j], style='List Number')

    def generate_contribution_list(self, df, document, initials):
        flatten = lambda l: [item for sublist in l for item in sublist]

        roles = df[self.role_tag].apply(lambda x: re.split(r"\s*[;]\s*", x))
        roles_set = set(flatten(roles.to_list()))

        initials_group_by_roles = {
            tag: [] for tag in roles_set
        }

        for idx in roles.index:
            initial = initials[idx]
            for tag in roles[idx]:
                initials_group_by_roles[tag].append(initial)

        document.add_heading("Author Contributions (in alphabetical order)", level=2)

        role_set_order = self.roles_priority.copy()
        role_set_order.extend(roles_set - set(role_set_order))

        for role_name in roles_set:
            paragraph = document.add_paragraph("")
            title = paragraph.add_run(role_name)
            title.italic = True
            paragraph.add_run("\n")

            author_list = initials_group_by_roles[role_name].copy()
            author_list.sort()

            for i, author in enumerate(author_list):
                if i > 0:
                    paragraph.add_run(", ")

                paragraph.add_run(author)


if __name__ == '__main__':
    df = pd.read_csv("enigmaPDtestwithROLES.csv")
    generator = InitialsGenerator()

    df_after = generator.transform(df)
    df_after.to_csv("initials_output.csv")


